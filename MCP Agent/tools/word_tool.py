"""
Word文档生成工具 - 使用 python-docx 创建Office Word文档
"""
import sys
import os
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from mcp_framework import tool
from typing import Dict, Any, List, Optional
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import re

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
OUTPUT_DIR = os.path.join(ROOT, 'outputs')


def clean_filename(filename: str) -> str:
    """清理文件名中的非法字符"""
    return re.sub(r'[<>:"/\\|?*]', '_', filename)


def parse_content_to_doc(doc: Document, content: str) -> None:
    """
    解析文本内容并添加到Word文档
    支持简单的Markdown风格语法：
    - # 标题
    - ## 二级标题
    - ### 三级标题
    - **粗体**
    - *斜体*
    - - 列表项
    - 1. 有序列表
    """
    lines = content.split('\n')
    in_list = False
    list_level = 0
    
    for line in lines:
        line = line.rstrip()
        
        if not line:
            doc.add_paragraph()
            continue
        
        # 标题
        if line.startswith('# '):
            heading = doc.add_heading(line[2:], level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            continue
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
            continue
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
            continue
        
        # 无序列表
        if line.startswith('- ') or line.startswith('* '):
            p = doc.add_paragraph(line[2:], style='List Bullet')
            continue
        
        # 有序列表
        if re.match(r'^\d+\. ', line):
            p = doc.add_paragraph(re.sub(r'^\d+\. ', '', line), style='List Number')
            continue
        
        # 普通段落
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.size = Pt(12)
        
        # 简单的格式处理
        # 粗体 **text**
        parts = re.split(r'(\*\*[^*]+\*\*)', line)
        p = doc.add_paragraph()
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = p.add_run(part[2:-2])
                run.bold = True
            else:
                p.add_run(part)


@tool("create_word_document")
def create_word_document(params: Dict[str, Any]) -> Dict[str, Any]:
    """
    创建Word文档工具
    
    参数:
        filename: 文件名（可选，默认使用时间戳）
        title: 文档标题（可选）
        content: 文档内容（支持简单Markdown风格语法）
        sections: 分节内容列表（可选，用于结构化文档）
            [{"title": "节标题", "content": "节内容"}]
    
    返回:
        Word文档创建结果
    """
    try:
        filename = params.get("filename", "document")
        title = params.get("title", "")
        content = params.get("content", "")
        sections = params.get("sections", [])
        
        # 确保输出目录存在
        os.makedirs(OUTPUT_DIR, exist_ok=True)
        
        # 清理文件名
        clean_name = clean_filename(filename)
        if not clean_name.endswith('.docx'):
            clean_name += '.docx'
        
        filepath = os.path.join(OUTPUT_DIR, clean_name)
        
        # 创建文档
        doc = Document()
        
        # 设置默认字体
        doc.styles['Normal'].font.name = '微软雅黑'
        doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        
        # 添加标题
        if title:
            title_para = doc.add_heading(title, level=0)
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_run = title_para.runs[0]
            title_run.font.size = Pt(18)
            title_run.font.bold = True
            title_run.font.color.rgb = RGBColor(0, 0, 0)
            doc.add_paragraph()
        
        # 添加内容
        if content:
            parse_content_to_doc(doc, content)
        
        # 添加分节内容
        if sections:
            for section_data in sections:
                sec_title = section_data.get("title", "")
                sec_content = section_data.get("content", "")
                
                if sec_title:
                    doc.add_heading(sec_title, level=2)
                
                if sec_content:
                    parse_content_to_doc(doc, sec_content)
        
        # 保存文档
        doc.save(filepath)
        
        return {
            "success": True,
            "filepath": filepath,
            "filename": clean_name,
            "message": f"Word文档已创建：{clean_name}",
            "location": OUTPUT_DIR
        }
        
    except ImportError:
        return {
            "success": False,
            "error": "需要安装 python-docx 库，请运行：pip install python-docx"
        }
    except Exception as e:
        return {
            "success": False,
            "error": f"创建Word文档时出错：{str(e)}"
        }


@tool("create_office_document")
def create_office_document(params: Dict[str, Any]) -> Dict[str, Any]:
    """
    通用办公文档创建工具（Word格式）
    
    这是一个用户友好的接口，自动根据内容类型创建合适的文档
    
    参数:
        document_type: 文档类型（"report", "letter", "memo", "plan", "general"）
        title: 文档标题
        content: 文档主要内容
        author: 作者（可选）
        date: 日期（可选，默认为今天）
    
    返回:
        文档创建结果
    """
    try:
        doc_type = params.get("document_type", "general").lower()
        title = params.get("title", "文档")
        content = params.get("content", "")
        author = params.get("author", "")
        date_str = params.get("date", "")
        
        # 根据文档类型生成文件名
        type_prefixes = {
            "report": "报告",
            "letter": "信函",
            "memo": "备忘录",
            "plan": "计划",
            "general": "文档"
        }
        prefix = type_prefixes.get(doc_type, "文档")
        filename = f"{prefix}_{title}"
        
        # 构建结构化内容
        full_content = ""
        if author:
            full_content += f"作者：{author}\n"
        if date_str:
            full_content += f"日期：{date_str}\n"
        if author or date_str:
            full_content += "\n"
        full_content += content
        
        return create_word_document({
            "filename": filename,
            "title": title,
            "content": full_content
        })
        
    except Exception as e:
        return {
            "success": False,
            "error": f"创建文档时出错：{str(e)}"
        }


if __name__ == "__main__":
    from mcp_framework import serve
    serve()
